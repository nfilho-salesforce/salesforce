#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gera a Ordem de Serviço (SOW) ZELLO/SERPRO-SWE em .docx.
Âncora de formato: OS Dataprev (Anexo F - Serviços Profissionais sob Demanda).
Escopo: Marketing Cloud (jornadas de comunicação) + Agentforce.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x1B, 0x5E, 0x8E)      # títulos de seção
HDR_FILL = "44546A"                     # cabeçalho de tabela (azul-cinza)
TOTAL_FILL = "1F3864"

doc = Document()

# ---------- estilos base ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def h1(num, text):
    p = doc.add_paragraph()
    r = p.add_run(f"{num}   {text}")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = BLUE
    p.space_before = Pt(14)
    # linha inferior
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "1B5E8E")
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def h2(num, text):
    p = doc.add_paragraph()
    r = p.add_run(f"{num}   {text}")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = BLUE
    return p

def h3(num, text):
    p = doc.add_paragraph()
    r = p.add_run(f"{num}   {text}")
    r.bold = True; r.font.size = Pt(11)
    return p

def sub(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(10.5)
    return p

def para(text, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p

def role(name, alcance, atividades):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(name + ": "); r.bold = True
    p.add_run(alcance)
    sp = doc.add_paragraph(); sp.paragraph_format.left_indent = Inches(0.5)
    rr = sp.add_run("Principais Atividades:"); rr.bold = True
    for a in atividades:
        b = doc.add_paragraph(style="List Bullet 2")
        b.add_run(a)

# ============================================================
# CABEÇALHO / TÍTULO
# ============================================================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Ordem de Serviço"); r.bold = True; r.font.size = Pt(20)
sf = doc.add_paragraph(); sf.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sf.add_run("salesforce  |  professional services")
rs.font.size = Pt(11); rs.font.color.rgb = BLUE; rs.bold = True
dr = doc.add_paragraph(); dr.alignment = WD_ALIGN_PARAGRAPH.CENTER
rd = dr.add_run("RASCUNHO / DRAFT — v1.0 — gerado em 26/08/2026")
rd.italic = True; rd.font.size = Pt(9); rd.font.color.rgb = RGBColor(0x88,0x88,0x88)

# ============================================================
# 1. DEFINIÇÕES
# ============================================================
h1("1.", "Definições")
para('Esta Ordem de Serviço, efetiva a partir da última data das assinaturas das partes abaixo '
     '("Data Efetiva da Ordem de Serviço") aplica-se e está sujeita aos termos do Anexo F - Serviços '
     'Profissionais sob Demanda - SOW entre "SFDC" e "Cliente" como definido abaixo. A SFDC se reserva '
     'o direito de rejeitar esta Ordem de Serviço se não for assinada até a Data de Validade fornecida abaixo.')

defs = [
    ('"SFDC"', "Salesforce Tecnologia Ltda."),
    ('"Cliente"', "SERVIÇO FEDERAL DE PROCESSAMENTO DE DADOS - SERPRO"),
    ("Endereço do Cliente", "[preencher endereço do Cliente]"),
    ("Gerente de Projetos do Cliente", "[nome / e-mail do PM do Cliente]"),
    ("SOW", "Anexo F - Serviços Profissionais sob Demanda - SOW"),
    ("Ordem de Serviço", "OS#[nº] — ZELLO/SERPRO-SWE"),
    ("Data de Validade da Ordem de Serviço", "[dd/mm/aaaa]"),
]
tb = doc.add_table(rows=len(defs), cols=2); tb.style = "Table Grid"
for i, (k, v) in enumerate(defs):
    c0 = tb.cell(i, 0); c0.paragraphs[0].add_run(k).bold = True
    tb.cell(i, 1).paragraphs[0].add_run(v)
tb.columns[0].width = Inches(2.4); tb.columns[1].width = Inches(4.1)

para("Na hipótese de, ao esgotarem as horas de Serviços Profissionais sob Demanda e/ou valores autorizados sob "
     "esta Ordem de Serviço, Serviços Profissionais sob Demanda adicionais sejam necessários para o progresso do "
     "projeto, as partes executarão um Pedido(s) de Alteração para continuidade dos Serviços Profissionais sob Demanda.")

# ============================================================
# 2. SUMÁRIO
# ============================================================
h1("2.", "Sumário da Ordem de Serviço")
rows = [
    ("Ordem de Serviço", "Horas Estimadas", "Valores Estimados de Serviços Profissionais"),
    ("OS#[nº] — ZELLO/SERPRO-SWE", "[TBD]", "R$ [TBD]"),
    ("Subtotal de Valores Estimados de Serviços Profissionais", "", "R$ [TBD]"),
    ("Taxa de imposto a partir da data de vigência", "", "6,55%"),
    ("Imposto Estimado", "", "R$ [TBD]"),
    ("Total Estimado com Impostos", "", "R$ [TBD]"),
]
st = doc.add_table(rows=len(rows), cols=3); st.style = "Table Grid"
for j in range(3):
    set_cell_bg(st.cell(0, j), HDR_FILL)
    rr = st.cell(0, j).paragraphs[0].add_run(rows[0][j]); rr.bold = True
    rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
for i in range(1, len(rows)):
    for j in range(3):
        st.cell(i, j).paragraphs[0].add_run(rows[i][j])

# ============================================================
# 3. DESCRIÇÃO
# ============================================================
h1("3.", "Descrição da Ordem de Serviço")
h2("3.1", "Escopo")
para("O Cliente solicitou a execução do seguinte escopo na Ordem de Serviço definida na Seção 1 acima. "
     "O escopo desta Ordem de Serviço compreende dois serviços de implementação Salesforce: (i) Agentforce "
     "e (ii) Marketing Cloud, este último voltado à orquestração de jornadas de comunicação com o cidadão/usuário.")

# ---- 3.1.1 Agentforce ----
h3("3.1.1", "Serviço de Implementação Agentforce")
para("Sujeito aos termos desta Ordem de Serviço e do Contrato firmado entre a SFDC e o Cliente, a SFDC fornecerá "
     "os Serviços Profissionais sob Demanda estabelecidos abaixo. A SFDC fornecerá serviços para implementar e "
     "auxiliar o Cliente a avaliar e revisar as capacidades de Agentforce atendendo aos casos de uso de Experiência "
     "do Agente mapeados com o Cliente na etapa de planejamento do projeto.")

sub("Descoberta")
for x in [
    "Conduzir workshops colaborativos para articular a visão estratégica e os objetivos do Cliente para futuras experiências de agentes (agentic experiences).",
    "Facilitar a avaliação dos sistemas, processos e arquitetura de experiência do estado atual relevantes para fluxos de trabalho habilitados por Agente.",
    "Apoiar a validação dos pré-requisitos organizacionais e técnicos necessários para permitir o desenvolvimento e implantação escaláveis de Agentes.",
    "Fornecer orientação na identificação e priorização de experiências de agentes alinhadas ao valor de negócio e considerações de viabilidade.",
    "Auxiliar no alinhamento das partes interessadas em uma visão compartilhada de experiência de agente e resultados de valor esperados.",
    "Revisar estimativas preliminares de ROI e direcionadores de valor para informar a estratégia de investimento e priorização de Agentes.",
]:
    bullet(x)

sub("Desenvolvimento e Implantação de Agente")
for x in [
    "Aconselhar sobre o projeto e especificação do Agente por meio de sessões de trabalho colaborativas para definir o comportamento, escopo e interações do usuário desejado.",
    "Configurar o Data Cloud e ingerir os dados necessários para o desenvolvimento do Agente.",
    "Configurar soluções de Agente (prompts/ações/tópicos/instruções/fluxos).",
    "Executar a construção e as iterações de teste do Agente, garantindo o alinhamento com os requisitos documentados e os objetivos de usabilidade.",
    "Facilitar revisões de trabalho regulares de construções de Agente em andamento com as partes interessadas do Cliente para incorporar feedback e refinar a funcionalidade.",
    "Auxiliar na validação da prontidão do ambiente para a implantação do Agente, incluindo a configuração do CRM e as dependências de acesso a dados.",
    "Aconselhar sobre as melhores práticas para a implantação do Agente, incluindo protocolos de teste e planejamento de ativação inicial.",
    "Apoiar a equipe do Cliente na implantação de um piloto inicial de Agente para um grupo de usuários segmentado para feedback e refinamento.",
    "Esboçar os próximos passos recomendados e critérios para expandir o uso do Agente e fazer a transição para uma implantação em produção mais ampla.",
]:
    bullet(x)

sub("Desenvolvimento e Apresentação de Roteiro")
for x in [
    "Facilitar a análise de requisitos e áreas de oportunidade para uma Organização de Agentes (Agentic Organization) em domínios de negócios priorizados.",
    "Apoiar avaliações de viabilidade e análise de ROI para informar a priorização de futuras experiências de Agentes.",
    "Conduzir sessões de revisão do roteiro inicial de experiência de agente com as partes interessadas do cliente para garantir o alinhamento de negócios.",
    "Alinhar com as partes interessadas do cliente em uma abordagem de seguimento (go-forward approach), incluindo propriedade e métricas de sucesso.",
    "Recomendar considerações consultivas para fases futuras, incluindo operações do Agente, otimização, governança e estratégia de escala.",
]:
    bullet(x)

sub("Pré-requisitos do Cliente")
for x in [
    "Compromisso de Negócios para Implantação do Agentforce: a equipe do Cliente tem patrocínio executivo e autoridade de negócios para implantar em produção.",
    "Compromisso Técnico para Implantação do Agentforce: a equipe do Cliente tem autoridade para resolver bloqueadores de implantação técnica e implantar em produção dentro do cronograma do engajamento.",
    "Desafios de Negócios Chave e Casos de Uso Validados: problema claramente definido a ser resolvido, com valor associado e alinhamento validado com o escopo e as premissas do Agentforce Fase 0.",
    "Equipe de Projeto do Cliente Estabelecida: o Cliente fornece uma equipe de projeto dedicada para colaboração ágil diária.",
    "Licença(s) e Créditos Salesforce Necessários: o Cliente provisionou as licenças e os créditos necessários para Core Cloud(s), Agentforce e Data Cloud.",
    "Disponibilidade de Dados e Ambientes Salesforce: um ambiente sandbox de CRM Salesforce atualizado, fornecido pelo cliente, com dados de produção reais está disponível para que a SFDC execute o trabalho esperado.",
]:
    bullet(x)

sub("Premissas de Escopo Agentforce")
for x in [
    "A org/ambiente Salesforce possui dados significativos de qualidade razoável.",
    "O Data Cloud está provisionado e pronto, incluindo o sandbox. A org inicial do Data Cloud está co-localizada com a org do Salesforce CRM.",
]:
    bullet(x)
bullet("Para cada caso de uso, o escopo pré-definido do Agente inclui:")
for x in [
    "1 experiência de Agente pré-definida para 1 marca com até 2 tópicos pré-definidos.",
    "Dados residentes no Org do Salesforce (até 4 objetos de 1 org, incluindo <20 campos por objeto).",
    "Para Agentes RAG, ingestão de apenas Salesforce Knowledge.",
    "Até 2 ações padrão.",
    "1 ação personalizada com dados residentes no Org do Salesforce (sem dados externos, sem autenticação).",
]:
    bullet(x, level=1)
for x in [
    "Espera-se que o trabalho ocorra nos ambientes sandbox/desenvolvimento do Salesforce do cliente.",
    "A implantação do Agente ocorrerá no Salesforce CRM.",
    "A implantação em canais externos será realizada pelo cliente, com o suporte da SFDC se a implantação for aprovada durante o cronograma do engajamento.",
    "Os recursos padrão serão usados sempre que for razoavelmente possível; a configuração cliques sem código será usada sempre que for razoavelmente possível; código personalizado será usado somente se a SFDC considerar necessário.",
    "O escopo assume um único idioma e uma única moeda, conforme aplicável.",
]:
    bullet(x)

sub("Exclusões de Escopo Agentforce")
bullet("Atividades não listadas nesta Ordem de Serviço são consideradas fora de escopo, incluindo o seguinte:")
for x in [
    "Governança e abordagem operacional.",
    "Migração de chat/agente em tempo real ou bots.",
    "Migração de dados para o Salesforce CRM.",
    "Gerenciamento ou melhorias da qualidade dos dados.",
]:
    bullet(x, level=1)
for x in [
    "A SFDC entregará o(s) Agente(s) em funcionamento. A implantação é, em última análise, a critério do Cliente.",
    "Qualquer trabalho não especificado nesta SOW é considerado fora de escopo, incluindo, sem limitação: migração de dados legados, APIs personalizadas, configuração de aplicativo móvel, serviços de gerenciamento de alterações, teste e otimização de performance, requisitos não funcionais (conformidade de segurança, padrões de acessibilidade, certificações regulatórias), gestão de recursos do cliente e suporte contínuo/manutenção/melhorias futuras.",
]:
    bullet(x)

# ---- 3.1.2 Marketing Cloud ----
h3("3.1.2", "Serviço de Implementação Marketing Cloud")
para("Um dos objetivos do projeto com esta OS é habilitar a comunicação proativa e transacional com o "
     "cidadão/usuário, integrando as jornadas de notificação (WhatsApp/E-mail/SMS) aos eventos de vida ou "
     "serviços/políticas públicas do Cliente. A SFDC fornecerá serviços para desenhar e configurar as jornadas "
     "de comunicação na plataforma Marketing Cloud, incluindo a orquestração e a integração com o Agentforce "
     "para transbordo de atendimento.")

sub("Orquestração de Jornadas de Comunicação (Journey Builder)")
for x in [
    "Desenvolvimento de réguas de relacionamento automatizadas para guiar o cidadão/usuário através dos serviços e comunicações.",
    "Jornadas de Notificação Ativa e transacional pelos canais WhatsApp/E-mail/SMS.",
    "Integração Bidirecional: configuração de fluxos no Journey Builder que iniciam no WhatsApp (notificação ativa) e transbordam para o Chatbot/Agentforce em caso de resposta do cidadão, mantendo o contexto da conversa.",
    "Definição e configuração de gatilhos por evento (event-driven) a partir das fontes de dados/eventos do Cliente.",
]:
    bullet(x)

sub("Integração de Dados")
for x in [
    "Conexão com as fontes de dados para garantir que a comunicação seja relevante e personalizada.",
    "Sincronização com Data Cloud/Salesforce ou CDP externo.",
    "Carga de segmentação para a jornada no Marketing Cloud; a estratégia será definida com o Cliente em tempo de projeto.",
]:
    bullet(x)

sub("Premissas do Escopo Marketing Cloud")
for x in [
    "Aprovação de Templates (WhatsApp): o Cliente é responsável pela submissão e aprovação dos templates de mensagens (HSMs) junto à Meta (WhatsApp), processo necessário para notificações ativas.",
    "Conteúdo e Criativos: o Cliente fornecerá todos os textos, imagens e diretrizes de tom de voz (brand guidelines) para a criação de mensagens. A equipe de Serviços Profissionais SFDC focará na configuração técnica e funcional das jornadas.",
    "Bases de Dados: as bases de audiência (listas de contatos) devem estar higienizadas e disponíveis no ambiente Salesforce (ou Data Cloud) com os devidos campos de chave primária para a ativação das jornadas.",
    "O escopo assume um único idioma e uma única moeda, conforme aplicável; recursos padrão e configuração sem código serão priorizados.",
]:
    bullet(x)

sub("Exclusões do Escopo Marketing Cloud")
for x in [
    "Criação de conteúdo, peças criativas ou design gráfico das mensagens.",
    "Aprovação e gestão de templates/HSM junto à Meta (WhatsApp).",
    "Desenvolvimento de pipelines de ingestão ou transformação de dados (ETL/ELT) não previstos nesta SOW.",
    "Saneamento e garantia da qualidade das bases de audiência de origem.",
    "Integração com canais ou provedores de mensageria não mapeados na fase de desenho.",
]:
    bullet(x)

# ---- 3.1.3 Evidência ----
h3("3.1.3", "Evidência e Comprovação de Entrega")
para("Caberá à equipe de Serviços Profissionais da Salesforce, através do Gerente de Projetos, a elaboração e "
     "apresentação das evidências técnicas que comprovem a plena operabilidade de seus entregáveis. Como evidência "
     "técnica será considerada a sequência de telas que comprovam o funcionamento das jornadas e do agente, e uma "
     "etapa de demonstração ao requisitante.")
para("Adicionalmente, é mandatória a entrega dos artefatos documentais (incluindo, mas não se limitando a, "
     "relatórios de testes e atas de homologação) referentes às funcionalidades mapeadas e entregues dentro do "
     "contexto e alcance do escopo descritos nesta Ordem de Serviço. A validação formal destas comprovações constitui "
     "condição para o processamento do Aceite Final da Ordem de Serviço.")

# ---- 3.1.4 Premissas Gerais ----
h3("3.1.4", "Premissas Gerais")
para("Para garantir o sucesso e o cumprimento do cronograma deste SOW, aplicam-se as seguintes premissas:")
sub("3.1.4.1  Premissas e Responsabilidades Gerais")
for x in [
    "O início de cada sprint de desenvolvimento está condicionado à entrega da documentação técnica e acessos por parte do Cliente.",
    "Aceite via Dados Sintéticos: o aceite técnico e faturamento das etapas de IA (Agentforce) podem ser feitos via Mock APIs ou dados simulados, caso a real integração sofra atrasos por parte do cliente ou de terceiros.",
    "Governança Executiva: estabelecer um comitê quinzenal de alto nível para destravar impedimentos políticos ou técnicos que fujam da alçada do time de projeto.",
    "O escopo pode ser limitado pela complexidade dos requisitos técnicos e modelos de dados do Cliente, bem como pela funcionalidade e recursos da Aplicação.",
    "Os recursos padrão serão usados sempre que for razoavelmente possível; a configuração cliques sem código será usada sempre que possível; código personalizado será usado somente se a SFDC considerar necessário.",
    "O escopo assume um único idioma e uma única moeda, conforme aplicável.",
    "Licença(s) e Créditos Salesforce Necessários: o Cliente provisionou as licenças e os créditos necessários para as Clouds envolvidas na solução (Marketing Cloud, Agentforce e Data Cloud).",
    "Público-Alvo Técnico: as sessões de transferência de conhecimento são destinadas exclusivamente a Administradores, Desenvolvedores e Arquitetos do Cliente. Não está incluído treinamento para usuários finais.",
    "Modelo \"Train the Trainer\": a SFDC capacitará os multiplicadores (treinadores) do Cliente. O Cliente é responsável por replicar esse treinamento para suas equipes internas e usuários finais.",
    "Idioma: os Serviços Profissionais e as interfaces serão configurados exclusivamente em Português (Brasil).",
    "Virtualidade: os Serviços Profissionais serão entregues de forma 100% virtual/remota, salvo acordo prévio por escrito.",
    "Produtos de Terceiros: a SFDC não se responsabiliza pelo suporte técnico ou manutenção de plataformas de terceiros, como a API oficial do WhatsApp (Meta), caso seja requerida.",
]:
    bullet(x)
sub("3.1.4.2  Exclusões do Escopo Gerais")
para("As seguintes atividades não estão incluídas neste SOW/OS e exigirão um Pedido de Mudança (Change Order) se solicitadas:")
for x in [
    "Configuração de aplicativo móvel.",
    "Mudança Centrada no Ser Humano ou serviços educacionais.",
    "Serviços de gerenciamento de alterações.",
    "Migração de Dados: nenhuma migração de dados legados.",
    "Execução de treinamentos presenciais ou em sala de aula para grandes grupos ou usuários finais.",
    "Suporte operacional continuado (\"Sustentação\" ou \"AMS\") após o término do período de Handover (Hypercare não contratado nesta OS).",
    "Atrasos na implantação causados pelo Cliente não serão responsabilidade da SFDC.",
    "Qualquer trabalho não especificado nesta OS é considerado fora de escopo.",
]:
    bullet(x)

# ============================================================
# 3.2 VALORES ESTIMADOS
# ============================================================
h2("3.2", "Valores Estimados de Serviços")
para("As horas nesta proposta foram estimadas com base nas informações disponíveis no momento de sua elaboração. "
     "Durante a fase inicial de planejamento e desenho da solução, essas estimativas serão detalhadas e validadas "
     "em conjunto entre as partes. Caso sejam identificadas necessidades de ajuste em escopo ou esforço, as revisões "
     "correspondentes serão discutidas e formalmente acordadas antes de iniciar a implementação.")
para("Os valores dos Serviços Profissionais sob Demanda para execução desta Ordem de Serviço estão especificados abaixo:")

res_hdr = ["Recursos", "Preço por Hora\n(sem tributos)", "Horas\nEstimadas",
           "Preço por Serviços\nProfissionais", "Valor Consolidado\nde Tributos", "Valor Bruto\nEstimado"]
recursos = [
    "Project Manager",
    "Solution Architect - Marketing Cloud",
    "Technical Architect - Marketing Cloud",
    "Technical Consultant - Marketing Cloud",
    "Technical Architect - Agentforce",
    "Technical Consultant - Agentforce",
]
rt = doc.add_table(rows=len(recursos) + 2, cols=6); rt.style = "Table Grid"
for j, htxt in enumerate(res_hdr):
    set_cell_bg(rt.cell(0, j), HDR_FILL)
    rr = rt.cell(0, j).paragraphs[0].add_run(htxt); rr.bold = True
    rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); rr.font.size = Pt(9)
for i, nome in enumerate(recursos, start=1):
    rt.cell(i, 0).paragraphs[0].add_run(nome)
    for j in range(1, 6):
        rt.cell(i, j).paragraphs[0].add_run("[TBD]")
last = len(recursos) + 1
for j in range(6):
    set_cell_bg(rt.cell(last, j), TOTAL_FILL)
    txt = "TOTAL" if j == 0 else ("-" if j == 1 else "R$ [TBD]" if j >= 3 else "[TBD]")
    rr = rt.cell(last, j).paragraphs[0].add_run(txt); rr.bold = True
    rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

# ============================================================
# 3.3 DURAÇÃO
# ============================================================
h2("3.3", "Estimativas de duração dos Serviços Profissionais sob Demanda")
para("O projeto foi estimado com a duração de até [nº] semanas a partir do engajamento dos recursos e início do "
     "projeto. A tabela abaixo representa o esforço estimado em horas do projeto por tipo de Serviço Profissional "
     "Sob Demanda:")
dur_rows = [
    ("Tipo de Serviço", "Horas Estimadas"),
    ("Implementação de Agentes (Agentforce)", "[TBD]"),
    ("Implementação Marketing Cloud", "[TBD]"),
    ("TOTAL", "[TBD]"),
]
dt = doc.add_table(rows=len(dur_rows), cols=2); dt.style = "Table Grid"
for j in range(2):
    set_cell_bg(dt.cell(0, j), HDR_FILL)
    rr = dt.cell(0, j).paragraphs[0].add_run(dur_rows[0][j]); rr.bold = True
    rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
for i in range(1, len(dur_rows)):
    for j in range(2):
        cell = dt.cell(i, j)
        rr = cell.paragraphs[0].add_run(dur_rows[i][j])
        if dur_rows[i][0] == "TOTAL":
            set_cell_bg(cell, TOTAL_FILL); rr.bold = True
            rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

# ============================================================
# 3.4 ESCOPO DE ATUAÇÃO DOS RECURSOS
# ============================================================
h2("3.4", "Escopo de atuação dos Recursos ou Perfis Técnicos Salesforce")
para("Segue a descrição dos perfis técnicos de Serviços Profissionais Salesforce presentes neste projeto, com o "
     "alcance de atuação e as principais atividades por perfil no contexto das duas Clouds em escopo — Marketing "
     "Cloud (jornadas de comunicação) e Agentforce.")

role("Project Manager (Gerente de Projetos)",
     "Responsável pela gestão ponta a ponta desta Ordem de Serviço, garantindo a entrega dentro do escopo, prazo e "
     "qualidade acordados. Atua como principal ponto de contato entre a equipe de Serviços Profissionais Salesforce "
     "e o Cliente, e como orquestrador dos recursos de Marketing Cloud e Agentforce.",
     ["Planejamento e condução do projeto sob metodologia ágil (sprints), incluindo cerimônias, backlog e roadmap "
      "de entregas das jornadas de comunicação (Marketing Cloud) e dos agentes (Agentforce).",
      "Gestão de escopo, cronograma, riscos, dependências e Change Orders, mantendo o registro e a comunicação de status executivo (comitê quinzenal).",
      "Coordenação dos recursos técnicos (arquitetos e consultores de MC e Agentforce) e articulação com os pontos focais de negócio e TI do Cliente.",
      "Governança das premissas e pré-requisitos: aprovação de templates WhatsApp/HSM, provisionamento de licenças/créditos, disponibilidade de dados e ambientes.",
      "Condução das aceitações (homologação/UAT), consolidação das evidências de entrega e processamento do Aceite Final da Ordem de Serviço."])

role("Solution Architect - Marketing Cloud",
     "Atua como elo entre os objetivos de comunicação e relacionamento do Cliente e o desenho funcional da solução "
     "de Marketing Cloud, garantindo que as jornadas de comunicação entreguem valor de negócio e boa experiência "
     "ao destinatário.",
     ["Condução da descoberta e do desenho funcional das jornadas de comunicação (réguas de relacionamento, "
      "notificações ativas e transacionais) no Journey Builder.",
      "Definição da estratégia de segmentação, dos canais (WhatsApp/E-mail/SMS) e da orquestração entre notificação "
      "ativa e transbordo para atendimento/Agentforce, preservando o contexto da conversa.",
      "Modelagem do data model de marketing (entradas, atributos, chaves) e alinhamento com as fontes de dados/Data "
      "Cloud para relevância e personalização das mensagens.",
      "Definição das métricas de sucesso das jornadas (entrega, abertura, resposta, conversão) e dos critérios de aceite funcionais.",
      "Tradução dos requisitos de negócio em backlog de configuração e orientação ao Technical Architect e ao Technical Consultant de Marketing Cloud."])

role("Technical Architect - Marketing Cloud",
     "Autoridade técnica da plataforma Marketing Cloud, responsável pela arquitetura, integrações, segurança e "
     "performance das jornadas de comunicação.",
     ["Definição da arquitetura técnica de Marketing Cloud e do modelo de integração de dados (Data Cloud/CRM/CDP "
      "externo) para ativação das jornadas.",
      "Desenho dos fluxos de integração bidirecional (entrada via WhatsApp → Journey Builder → transbordo para "
      "Agentforce), garantindo a manutenção de contexto da conversa.",
      "Definição de padrões de segurança, privacidade (LGPD) e conformidade no tratamento das bases de audiência e chaves primárias.",
      "Governança da configuração de canais, do provisionamento de templates/HSM e dos requisitos não funcionais "
      "(volumetria de disparos, janelas de envio, entregabilidade).",
      "Suporte técnico ao Technical Consultant de MC e mitigação de riscos de integração e de entregabilidade."])

role("Technical Consultant - Marketing Cloud",
     "Responsável pela materialização técnica das jornadas de comunicação, atuando na configuração direta da "
     "plataforma Marketing Cloud conforme o desenho funcional e arquitetural.",
     ["Configuração de jornadas no Journey Builder, réguas de relacionamento, automações e Data Extensions.",
      "Configuração dos canais (WhatsApp/E-mail/SMS), integração dos templates/HSM aprovados e das entradas de dados/segmentação.",
      "Implementação da integração bidirecional com o Agentforce/Chatbot para transbordo de atendimento e retorno de contexto.",
      "Execução de testes unitários e de integração das jornadas, validação de entregabilidade e apoio na homologação (UAT).",
      "Documentação técnica das jornadas de comunicação implementadas."])

role("Technical Architect - Agentforce",
     "Autoridade técnica da solução Agentforce, responsável pela arquitetura do agente, Data Cloud, ações/integrações "
     "e requisitos não funcionais.",
     ["Definição da arquitetura do Agente (tópicos, instruções, ações, fluxos) e do modelo de dados (objetos, campos, "
      "RAG via Salesforce Knowledge).",
      "Desenho da configuração e da ingestão no Data Cloud necessária ao funcionamento do agente.",
      "Definição de padrões de segurança, privacidade e governança de prompts/ações, incluindo dependências de acesso a dados.",
      "Definição de requisitos não funcionais (tempo de resposta, prontidão do ambiente) e critérios de ativação do piloto.",
      "Suporte técnico ao Technical Consultant de Agentforce e mitigação de riscos técnicos complexos."])

role("Technical Consultant - Agentforce",
     "Responsável pela materialização técnica do Agente, atuando na construção, configuração e nas iterações de teste "
     "conforme o desenho acordado.",
     ["Configuração de soluções de Agente (prompts/ações/tópicos/instruções/fluxos) e configuração do Data Cloud/ingestão de dados.",
      "Construção e iteração das capacidades do agente, garantindo alinhamento aos requisitos documentados e aos objetivos de usabilidade.",
      "Implementação de ação(ões) padrão e personalizada(s) com dados residentes no Org do Salesforce.",
      "Execução de testes, apoio na validação de prontidão do ambiente e na implantação do piloto para grupo de usuários segmentado.",
      "Documentação técnica e apoio ao roadmap de expansão do uso do agente."])

# ============================================================
# 4. PAGAMENTO
# ============================================================
h1("4.", "Pagamento")
p = doc.add_paragraph()
p.add_run("É necessário um Pedido de Compra (\"PO\") para a compra ou pagamento dos produtos neste SOW? ")
p.add_run("Não.").bold = True
para("O Cliente não exige um PO para a compra ou para o pagamento dos Serviços Profissionais nesta Ordem de Serviço. "
     "Todos os Honorários são cobrados por horas efetivamente trabalhadas, que serão faturadas mensalmente e "
     "pagáveis de acordo com os termos do Contrato.")

# ============================================================
# 5. ASSINATURAS
# ============================================================
h1("5.", "Assinaturas")
para("Esta Ordem de Serviço pode ser assinada em contrapartidas, cada uma das quais será considerada como original. "
     "EM TESTEMUNHO DO QUE, as partes fizeram com que esta Ordem de Serviço fosse executada por seus representantes "
     "devidamente autorizados, conforme identificado abaixo.")
sig = doc.add_table(rows=4, cols=4); sig.style = "Table Grid"
set_cell_bg(sig.cell(0, 0), HDR_FILL); set_cell_bg(sig.cell(0, 2), HDR_FILL)
sig.cell(0, 0).merge(sig.cell(0, 1)); sig.cell(0, 2).merge(sig.cell(0, 3))
c = sig.cell(0, 0).paragraphs[0].add_run("Cliente"); c.bold = True; c.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
c = sig.cell(0, 2).paragraphs[0].add_run("SFDC"); c.bold = True; c.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
labels = ["Assinatura de Autorização", "Nome", "Título", "Data"]
for i, lbl in enumerate(labels[:3], start=1):
    sig.cell(i, 0).paragraphs[0].add_run(lbl).bold = True
    sig.cell(i, 2).paragraphs[0].add_run(lbl).bold = True

out = "/Users/nfilho/claude/Scopezilla/ZELLO-SERPRO-SWE/SOW_ZELLO-SERPRO-SWE_v1.0.docx"
doc.save(out)
print("Salvo:", out)

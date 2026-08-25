#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gerador de SOW — CLARO Brasil | CPQ Design (Communications Cloud)
Template: Salesforce PS — SOW Standard Language Library (SLL PT), Seções 1–7
Modelo de honorários: Fixed Fee (escopo fechado)
Idioma: PT-BR
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Paleta / helpers ----------
SF_BLUE = RGBColor(0x03, 0x2D, 0x60)      # azul institucional
SF_ACCENT = RGBColor(0x00, 0x70, 0xD2)    # azul Salesforce
GREY = RGBColor(0x54, 0x69, 0x8D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# Fonte base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = SF_BLUE
    r.font.size = Pt(15)
    return p

def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = SF_ACCENT
    r.font.size = Pt(12.5)
    return p

def h3(text):
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    r.font.color.rgb = GREY
    r.font.size = Pt(11)
    return p

def para(text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25*level)
    p.add_run(text)
    return p

def make_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
        set_cell_bg(hdr[i], '032D60')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return t

# ============================================================
# CAPA
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('STATEMENT OF WORK (SOW)')
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = SF_BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Salesforce Professional Services')
r.font.size = Pt(13); r.font.color.rgb = SF_ACCENT

doc.add_paragraph()
proj = doc.add_paragraph()
proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = proj.add_run('CLARO BRASIL — CPQ Design')
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = SF_BLUE
proj2 = doc.add_paragraph()
proj2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = proj2.add_run('Desenho de Catálogo Unificado e Estrutura de CPQ\nSalesforce Communications Cloud (EPC · BRE · Order Management)')
r.font.size = Pt(12); r.font.color.rgb = GREY

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('Modelo de Engajamento: Fixed Fee — Baseado em Outcomes (4 Outcomes)\nEntregas: 4 outcomes em cadência quinzenal · Fluxo de caixa em ~2 meses\nRegião: LATAM — Brasil')
r.font.size = Pt(11); r.italic = True; r.font.color.rgb = GREY

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ver.add_run('Versão 1.0 (Rascunho) · Confidencial')
r.font.size = Pt(9); r.font.color.rgb = GREY

doc.add_page_break()

# ============================================================
# SEÇÃO 1 — VISÃO GERAL
# ============================================================
h1('1. Visão Geral')

h2('1.1 Introdução')
para('Este Statement of Work ("SOW") descreve os serviços profissionais ("Serviços") a serem '
     'prestados pela Salesforce ("SFDC") ao Cliente ("Claro Brasil"), sob os termos do Master '
     'Subscription Agreement (ou contrato equivalente vigente) firmado entre as partes. O objetivo '
     'do engajamento é o desenho da estrutura de Catálogo Unificado e da lógica de CPQ na '
     'Salesforce Communications Cloud, traduzindo as regras de negócio, ofertas, planos e serviços '
     'técnicos da Claro Brasil em um modelo de dados estruturado que habilite venda ágil, cross-sell '
     'e up-sell, substituindo matrizes customizadas legadas por capacidades nativas (EPC/BRE).')

h2('1.2 Dados do Cliente')
make_table(
    ['Campo', 'Informação'],
    [
        ['Cliente', 'Claro Brasil'],
        ['Responsável por Contratos', '[A preencher]'],
        ['Contato de Faturamento', '[A preencher]'],
        ['Contato de Envio / Projeto', '[A preencher]'],
        ['CNPJ', '[A preencher]'],
    ],
    col_widths=[2.2, 4.3]
)

h2('1.3 Dados SFDC')
make_table(
    ['Campo', 'Informação'],
    [
        ['Entidade Legal SFDC', 'Salesforce (entidade contratante LATAM)'],
        ['Contato de Vendas', 'Nelson Stebulaitis Filho — Services Sales Solution Manager'],
        ['Região', 'LATAM — Brasil'],
    ],
    col_widths=[2.2, 4.3]
)

h2('1.4 Definições')
defs = [
    ('Contrato', 'significa o Contrato de Serviços Profissionais disponível em '
     'https://www.salesforce.com/company/legal/agreements.'),
    ('Aplicação', 'significa qualquer serviço online, baseado na web e componentes on-premises '
     'offline disponibilizados pela SFDC ou qualquer Afiliada da SFDC ao Cliente sob um contrato '
     'separado em vigor entre o Cliente e a SFDC ou qualquer Afiliada da SFDC.'),
    ('Pedido de Alteração', 'é uma emenda formal a esta SOW gerada pela SFDC e assinada por ambas as partes.'),
    ('Data de Vigência', 'será a última das respectivas datas de assinaturas desta SOW pelas partes e do '
     'recebimento da PO, se exigido pelo Cliente ou pela SFDC.'),
    ('Honorários', 'significam os honorários de Serviços Profissionais de acordo com o desempenho desta SOW.'),
    ('Prazo da SOW', 'significa a duração do desempenho dos Serviços Profissionais de acordo com esta SOW e '
     'qualquer Pedido de Alteração.'),
    ('T&E', 'significa todas as despesas relacionadas com viagens e outras despesas, incluindo, mas não se '
     'limitando ao transporte, quilometragem e despesas razoáveis desembolsadas, incorridas pela SFDC na '
     'execução dos Serviços Profissionais.'),
]
for term, d in defs:
    p = doc.add_paragraph()
    r = p.add_run(f'"{term}" '); r.bold = True; r.font.color.rgb = SF_BLUE
    p.add_run(d)

doc.add_page_break()

# ============================================================
# SEÇÃO 2 — ESCOPO
# ============================================================
h1('2. Escopo')

h2('2.1 Histórico e Objetivo')
para('A Claro Brasil, maior operadora de telecomunicações do país, opera cinco linhas de negócio '
     '(Móvel, TV, Banda Larga, Fixo e Aparelhos) com catálogos e regras de precificação mantidos em '
     'matrizes customizadas legadas. Esta fragmentação eleva o time-to-market (TTM) de novas ofertas '
     'e dificulta a venda de combos multi-LOB. Este engajamento — de natureza exclusivamente de '
     'DESIGN (não implementação produtiva) — visa desenhar, na Communications Cloud, um modelo de '
     'catálogo unificado de duas camadas (Produto Comercial + Produto Técnico/CFS), regras de CPQ '
     '(elegibilidade, compatibilidade e precificação via BRE) e o padrão de decomposição de pedidos '
     'para os sistemas legados, validado por uma Prova de Conceito (PoC) funcional. Esta seção 2.1 '
     'é informativa e descreve os objetivos de negócio em alto nível.')

h3('Objetivos de Negócio (KPIs-alvo)')
for k in [
    'Redução de 45% no TTM de lançamento de ofertas (horizonte de 12 meses).',
    '100% das ofertas (até 20 por LOB) mapeadas no EPC ao final da fase de design.',
    '100% das regras de precificação modeladas via BRE — zero matrizes customizadas.',
    'Erro de payload de Order Management (OM) → 0% na PoC.',
    'Resposta de elegibilidade por CEP em menos de 2 segundos.',
    'Redução de 30% no abandono de carrinho por incompatibilidade.',
    'Aumento de 15% no ticket médio por combo multi-LOB (6 meses pós go-live).',
]:
    bullet(k)

h2('2.2 Serviços de Implementação — Communications Cloud (Design)')
para('Os Serviços a seguir compõem a fase de Design da Communications Cloud (SFI — Industries / '
     'Communications Cloud Expert Services, ref. SLL 1.2.21 / 2.3.10). Todos os entregáveis são de '
     'natureza de desenho e prova de conceito.', italic=True)

h3('2.2.1 Modelagem do Catálogo de Produtos (EPC — TM Forum SID)')
for x in [
    'Modelagem das 5 Linhas de Negócio (Móvel, TV, Banda Larga, Fixo e Aparelhos).',
    'Até ~100 produtos comerciais (até 20 por LOB) mais SVAs e adicionais.',
    'Duas camadas: Commercial Product (o que o cliente compra) e Technical Product / CFS (o que é ativado na rede).',
    'Definição do melhor product template (Object Type vs. Product Specification).',
    'Revisão da estrutura de Catalogs para melhor experiência de prateleira.',
    'Até 5 promoções/descontos avaliados (Fidelidade, sazonais, promoção por Node).',
]:
    bullet(x)

h3('2.2.2 Atributos e Cardinalidade')
for x in [
    'Mapeamento de características por produto (velocidade, franquia, minutos).',
    'Estrutura de obrigatório / opcional / padrão (ex.: decodificador obrigatório no plano TV).',
    'Uso de Product e Offering Specification conforme o padrão TM Forum.',
]:
    bullet(x)

h3('2.2.3 Regras de CPQ — Elegibilidade, Compatibilidade e Precificação')
for x in [
    'Elegibilidade: quem pode comprar (ex.: Fibra 1GB apenas para CEP com viabilidade técnica).',
    'Compatibilidade: bloqueios e dependências cruzadas (planos antigos × nova tecnologia).',
    'Precificação: recorrente (mensalidade) + não-recorrente (instalação), fidelidade 12 meses, descontos escalonados/promocionais.',
    'Business Rules Engine (BRE) / Expression Sets substituindo as matrizes customizadas legadas de cálculo de preços.',
]:
    bullet(x)

h3('2.2.4 Decomposição de Pedidos (Order Management)')
for x in [
    'Desenho de como o produto comercial vendido se decompõe em ordens técnicas para os sistemas legados.',
    'Ex.: combo Internet + TV → ordem de campo (instalar modem) + ordem lógica (liberar sinal de TV).',
]:
    bullet(x)
p = doc.add_paragraph()
r = p.add_run('Nota: O redesenho do fluxo completo de orquestração está FORA do escopo (ver 2.4.2).')
r.italic = True; r.font.color.rgb = GREY

h3('2.2.5 Prova de Conceito (PoC) Funcional')
for x in [
    '1 oferta por linha de negócio configurada e testada no CPQ nativo (5 PoCs).',
    'Inclui: prateleira, carrinho e decomposição de pedido.',
    'Pré-requisito: Sandbox funcional provisionada pela Claro (ver 2.4.1).',
]:
    bullet(x)

h3('2.2.6 Knowledge Transfer (KT)')
bullet('Workshops para as equipes de TI e de Negócios apresentando o modelo proposto e a PoC.')

h2('2.3 Outcomes do Engajamento (Base de Entrega)')
para('Este engajamento é estruturado e faturado com base em 4 (quatro) outcomes, entregues em '
     'cadência quinzenal. Cada outcome possui critério de aceite objetivo; o aceite formal de cada '
     'outcome dispara a parcela correspondente (ver Seção 4). Este é o núcleo do modelo Fixed Fee '
     'baseado em outcomes.', bold=False)
make_table(
    ['#', 'Outcome (entrega)', 'Fase', 'Descrição', 'Critério de Aceite'],
    [
        ['1', 'Detalhamento Funcional das Ofertas', 'F1 — Imersão',
         'Documentação dos workshops de Discovery: detalhamento funcional das ofertas/produtos existentes por LOB (com TI e Negócios).',
         'Documento de detalhamento funcional revisado e aceito por TI e Negócios.'],
        ['2', 'Product Model Blueprint (PMB)', 'F2 — Modelagem',
         'Diagrama visual completo da árvore de produtos (comercial + técnica/CFS) no EPC + especificação de elegibilidade, compatibilidade e precificação via Motor de Regras.',
         'PMB e recomendações de regras aprovados pelos Product Owners das 5 LOBs.'],
        ['3', 'Sandbox Configurada (PoC Funcional)', 'F3 — Build',
         '1 oferta por LOB com carrinho CPQ, validações e payloads de decomposição OM funcionais (prateleira + carrinho + decomposição).',
         'Demonstração da PoC validada em Sandbox (UAT assinado pela Claro).'],
        ['4', 'Relatório Final + KT', 'F4 — Transferência',
         'Relatório final do projeto, PoC funcional validada e gravações dos workshops de transferência de conhecimento (TI + Negócios).',
         'Entrega do Relatório Final e conclusão dos workshops de KT.'],
    ],
    col_widths=[0.3, 1.6, 1.1, 2.3, 1.6]
)

h2('2.4 Premissas e Exclusões de Escopo')

h3('2.4.1 Premissas de Escopo')
p = doc.add_paragraph()
r = p.add_run('Cláusula de IA (condição do prazo/valor de 8 semanas): '); r.bold = True; r.font.color.rgb = SF_BLUE
p.add_run('O prazo de 8 semanas e o valor correspondente estão condicionados ao consentimento formal '
          'e explícito da Claro Brasil quanto à aplicação e uso de ferramentas de IA pela SFDC na '
          'execução dos Serviços. Na ausência deste consentimento, aplica-se o cenário de 12 semanas '
          '(ver Seção 4.1).')
for x in [
    'O volume total de ofertas avaliadas não excederá 100 itens complexos (até 20 por LOB) e 5 promoções.',
    'A modelagem proposta, baseada em capacidades nativas declarativas, eliminará a necessidade de customizações em código.',
    'O barramento corporativo possui estabilidade e capacidade técnica para receber mensagens assíncronas de OM.',
    'O cronograma é orientado a entregáveis — extensões dependem de aprovação formal via Pedido de Mudança.',
    'As regras de precificação vigentes serão fornecidas em formato legível no início da Fase 2.',
    'A Claro Brasil provisionará uma Sandbox funcional da Communications Cloud antes do início da Fase 3 (PoC Build) — pré-requisito inegociável.',
    'A Claro disponibilizará Product Owners com autonomia de aprovação por LOB durante a Fase 1.',
    'A Claro fornecerá amostras das matrizes de elegibilidade/precificação atuais (planilhas ou diagramas lógicos).',
    'Documentação técnica (Swagger/OpenAPI/WSDL) das APIs de Billing, OSS/BSS e Georreferenciamento entregue no início da Fase 1; documentação não fornecida até o início da Imersão é risco de extensão de escopo.',
    'Os Serviços são executados prioritariamente de forma remota, salvo workshops presenciais previamente acordados (sujeitos a T&E).',
]:
    bullet(x)

h3('2.4.2 Exclusões de Escopo')
for x in [
    'Implementação em produção (go-live) — este SOW cobre exclusivamente a fase de Design e PoC.',
    'Redesenho do fluxo de orquestração pós-decomposição de Order Management.',
    'Customizações em código Apex / hooks em métodos nativos.',
    'PoC da modelagem sobre customizações no CPQ nativo ou legadas.',
    'Saneamento e carga massiva de dados históricos (Installed Base).',
    'Integração de canais de venda adicionais.',
    'Integrações produtivas com Billing e OSS/BSS (apenas o padrão de decomposição é desenhado).',
    'Serviços de gestão de aplicação (AMS) pós-projeto.',
    'Fornecimento de licenças Salesforce (objeto de cotação/Order Form separado).',
]:
    bullet(x)

doc.add_page_break()

# ============================================================
# SEÇÃO 3 — METODOLOGIA E CRONOGRAMA
# ============================================================
h1('3. Metodologia e Cronograma')

h2('3.1 Cronograma')
para('O engajamento é entregue em 4 outcomes com previsibilidade de entrega a cada 2 (duas) semanas, '
     'resultando em um horizonte de execução de aproximadamente 2 meses (~8 semanas). A cadência '
     'quinzenal define os marcos de aceite e o fluxo de caixa (ver Seção 4).')
make_table(
    ['Marco', 'Outcome (entrega da fase)', 'Prazo (a partir do kickoff)'],
    [
        ['M1', 'Outcome 1 — Detalhamento Funcional das Ofertas (F1 Imersão)', 'Semana 2'],
        ['M2', 'Outcome 2 — Product Model Blueprint (PMB) (F2 Modelagem)', 'Semana 4'],
        ['M3', 'Outcome 3 — Sandbox Configurada / PoC Funcional (F3 Build)', 'Semana 6'],
        ['M4', 'Outcome 4 — Relatório Final + KT (F4 Transferência)', 'Semana 8'],
    ],
    col_widths=[0.7, 3.9, 1.9]
)
para('As datas específicas serão confirmadas no kickoff. As atividades de discovery, modelagem, '
     'build de PoC e KT (metodologia PS, 3.2) são conduzidas dentro desta cadência quinzenal. '
     'Atrasos em premissas de responsabilidade do Cliente (Seção 6) podem impactar o cronograma e '
     'serão tratados via Pedido de Alteração.', italic=True, size=9.5)

h2('3.2 Metodologia dos Serviços de Implementação Salesforce')
para('A Salesforce utiliza a Metodologia de Serviços Profissionais Salesforce ("MSPS") para executar os '
     'Serviços de Implementação. A MSPS envolve quatro etapas: Definição (Define), Desenho (Design), '
     'Entrega (Deliver) e Implantação (Deploy). Essa metodologia é adaptável e ajustada para atender as '
     'necessidades únicas de cada cliente, incorporando ambos os princípios waterfall e agile. A estrutura '
     'escolhida depende das necessidades do Cliente, da natureza do engajamento e das recomendações da SFDC '
     'sobre melhores práticas. Durante as etapas de Definição e Desenho, a SFDC e o Cliente acordarão quanto '
     'à abordagem planejada, incluindo como o escopo será gerenciado.')
para('Como parte da MSPS: a SFDC e o Cliente documentarão a funcionalidade de Aplicação configurada '
     'desejada como uma lista de necessidades funcionais por tipo de usuário ("Histórias de Usuário"); o '
     'desenvolvimento será feito em uma série de incrementos regulares ("Sprints"); as revisões de '
     'desenvolvimento são realizadas mediante a conclusão de cada Sprint; e a priorização contínua da '
     'História de Usuário, conforme mutuamente acordado por escrito (e-mail suficiente) entre o Proprietário '
     'do Produto do Cliente e a SFDC, ditará o conteúdo final da Aplicação configurada. Como consequência, o '
     'escopo preciso não pode ser integralmente determinado no início dos Serviços de Implementação.')

h3('Etapas de Definição e Desenho')
para('Durante as etapas de Definição e Desenho, a SFDC inicia o projeto e colaborativamente estabelece as '
     'metas e objetivos de resultados de negócio do Cliente, define como o projeto será gerenciado e o plano '
     'de execução (escopo, recursos, cronograma, orçamento e quality assurance), e estabelece a arquitetura '
     'em alto nível necessária para iniciar a etapa de Entrega. A SFDC coordenará, de forma sequencial, as '
     'reuniões de Alinhamento do Time de Projeto, Kickoff e Alinhamento da Metodologia, Arquitetura Técnica '
     'e Análise de Integração e Preparação da Construção.')
para('Entregáveis: Documentos de Planejamento (PMP, Cronograma de Referência, Desenho da Solução em '
     'Alto-Nível, Garantia de Qualidade e Planos de Testes) e Documentos de Desenho (matriz de perfil de '
     'usuário, desenho do modelo de objeto, mapa de história e backlog do produto), que uma vez aprovados '
     'governam o escopo, cronograma, orçamento e entregáveis.', size=9.5)

h3('Etapa de Entrega')
para('Na etapa de Entrega, a SFDC configura, desenvolve e testa a Aplicação em uma série de Sprints baseados '
     'no backlog do produto das Histórias de Usuário e nos Documentos de Planejamento e Desenho. A SFDC '
     'realizará o teste unitário e criará casos de teste para validar a conformidade com o critério de '
     'aceitação. O Cliente é responsável por conduzir os Testes Integrados de Sistema (SIT) e os Testes de '
     'Aceitação de Usuário (UAT) e por fornecer feedback e aceitação, com apoio da SFDC. Entregáveis: '
     'Aplicação Configurada e Conclusão do UAT.')

h3('Etapa de Implantação')
para('Neste engajamento de Design/PoC, a etapa de Implantação consolida o design final e a validação da PoC '
     'em Sandbox. A migração para ambiente produtivo (go-live) está fora do escopo (ver 2.4.2). Entregáveis: '
     'Plano de Implantação (abordagem de referência), Documentação de Histórias de Usuário/Backlog do Produto '
     'e Documentação da Aplicação configurada.')

h2('3.3 Aceitação')
para('Após a conclusão de cada Entregável, a SFDC enviará uma cópia completa de tal Entregável ao Cliente e, '
     'a pedido, demonstrará sua funcionalidade. O Cliente é responsável por revisar e testar todos os '
     'Entregáveis de acordo com os critérios de aceitação e planos de teste mutuamente acordados por escrito. '
     'O Cliente fornecerá notificação por escrito de aceitação; a falha em rejeitar um Entregável será '
     'considerada aceitação. Caso o Cliente determine, em julgamento razoável e de boa-fé, que um Entregável '
     'não satisfaz os critérios de aceitação, deverá notificar a SFDC por escrito dentro de dez (10) dias '
     'úteis após o envio, especificando as deficiências em detalhes. A SFDC empregará esforços comercialmente '
     'razoáveis para corrigir e reapresentar o Entregável, e o Cliente terá novos dez (10) dias úteis para '
     'reavaliação. A aceitação de Histórias de Usuário é um processo iterativo: o Cliente as aceita '
     'formalmente via aplicativo de rastreamento ou por escrito (e-mail aceitável); não havendo aceite nem '
     'rejeição, serão consideradas aceitas no primeiro dos dez (10) dias úteis após a demonstração ou dois '
     '(2) dias úteis após o término da Sprint em que foram concluídas.')

doc.add_page_break()

# ============================================================
# SEÇÃO 4 — HONORÁRIOS
# ============================================================
h1('4. Honorários')

h2('4.1 Honorários — Taxa Fixa com Base em Marcos (Fixed Fee / Outcome-Based)')
para('Os Serviços Profissionais são prestados com base em uma taxa fixa (Fixed Fee), atrelada à '
     'entrega de 4 (quatro) marcos/outcomes — não T&M (Tempo & Materiais). Um Pedido de Mudança '
     'pode ser necessário se houver alterações no escopo do projeto, recursos, cronograma, marcos, '
     'entregas, complexidade ou designação de um novo gerente de projeto do Cliente. A SFDC faturará '
     'ao Cliente após a aceitação de cada marco. T&E e impostos efetivos e razoáveis, se aplicável, '
     'serão faturados conforme os termos do Contrato.')
para('Como as 4 entregas têm esforço e duração equivalentes (2 semanas cada), o valor é dividido em '
     '4 parcelas iguais (25% cada), gerando fluxo de caixa previsível em ~2 meses. Os valores abaixo '
     'refletem o cenário de 8 semanas com uso de IA (ver premissa da Cláusula de IA em 2.4.1).', bold=True)
make_table(
    ['Marco / Outcome', 'Critério de Aceite', 'Preço estimado (BRL)', 'Tributos estimados (BRL)', 'Valor Bruto (BRL)'],
    [
        ['M1 — Detalhamento Funcional das Ofertas', 'Documento funcional aceito por TI e Negócios', 'R$ 160.854,16', 'R$ 11.274,42', 'R$ 172.128,58'],
        ['M2 — Product Model Blueprint (PMB)', 'PMB + regras aprovados pelos POs das 5 LOBs', 'R$ 160.854,16', 'R$ 11.274,42', 'R$ 172.128,58'],
        ['M3 — Sandbox Configurada (PoC)', 'PoC validada em Sandbox (UAT assinado)', 'R$ 160.854,16', 'R$ 11.274,42', 'R$ 172.128,58'],
        ['M4 — Relatório Final + KT', 'Relatório Final entregue e KT concluído', 'R$ 160.854,16', 'R$ 11.274,43', 'R$ 172.128,59'],
        ['TOTAL', '', 'R$ 643.416,64', 'R$ 45.097,69', 'R$ 688.514,33'],
    ],
    col_widths=[1.9, 1.9, 1.1, 1.1, 1.1]
)
para('*A SFDC apresentará os requisitos desenvolvidos, e o Cliente confirmará a aceitação com base '
     'nos critérios estabelecidos em até 10 (dez) dias. Estimativa não vinculativa: valores finais '
     'confirmados na assinatura. Moeda: BRL. Fator fiscal aplicado: Preço estimado = Valor Bruto × '
     '0,9345 (Tributos = 6,55%).', italic=True, size=9.5)
para('Cenário alternativo (sem Cláusula de IA — 12 semanas): Valor Bruto total de R$ 1.032.770,47 '
     '(com impostos), conforme ROM. Neste caso, os marcos e parcelas serão reajustados no Pedido de '
     'Mudança correspondente.', italic=True, size=9.5)

h2('4.2 Despesas de Viagem e Estadia (T&E)')
para('Eventuais despesas de T&E (workshops presenciais) serão faturadas ao custo, mediante aprovação '
     'prévia por escrito do Cliente, e não estão incluídas no Fixed Fee.')

h2('4.3 Informações Fiscais do Cliente')
make_table(
    ['Campo', 'Informação'],
    [
        ['Razão Social', 'Claro Brasil — [A preencher]'],
        ['CNPJ', '[A preencher]'],
        ['Isenção Fiscal', '[A preencher, se aplicável]'],
        ['Moeda', 'BRL (R$)'],
        ['Tributos', 'Estimados em 6,55% sobre o valor bruto (fator 0,9345)'],
    ],
    col_widths=[2.2, 4.3]
)

h2('4.4 Requisitos de Ordem de Compra (PO)')
para('Caso o Cliente exija PO para faturamento, esta deverá ser fornecida antes do início dos '
     'Serviços. A ausência de PO não isenta o Cliente das obrigações de pagamento previstas neste SOW.')

doc.add_page_break()

# ============================================================
# SEÇÃO 5 — CONDIÇÕES DE SERVIÇO
# ============================================================
h1('5. Recursos da SFDC')

h2('5.1 Alterações de Pessoal/Horas')
para('Desde que o valor total não exceda o Fixed Fee estimado na Seção 4, um Pedido de Alteração não é '
     'necessário se o Gerente de Projeto da SFDC realocar esforço entre as funções para atender aos Serviços '
     'Profissionais e ao escopo definido na Seção 2.')

h2('5.2 Trabalhos Faturáveis')
para('Todos os trabalhos conduzidos por recursos sob esta SOW são considerados Serviços Profissionais '
     'faturáveis e não estão relacionados a serviços habilitados fornecidos por programas SFDC (por exemplo, '
     'Suporte ao Produto, Sucesso do Cliente etc.). O Cliente submeterá os casos de suporte técnico em geral '
     'à equipe de suporte da Salesforce associada ao uso dos serviços online do Cliente. A SFDC não executará '
     'tarefas gerais de suporte técnico sob esta SOW.')

h2('5.3 Proficiência')
para('A SFDC proverá proficiência exclusivamente na Aplicação e não garante que os recursos sejam '
     'proficientes em outras linguagens de programação ou plataformas.')

h2('5.4 Horário Fora do Expediente')
para('Os Serviços Profissionais serão realizados durante o horário comercial normal (de segunda a '
     'sexta-feira, das 8h30 às 17h30 no fuso horário do local onde está localizado o recurso de Serviços '
     'Profissionais), salvo acordo em contrário por escrito entre as partes (e-mail é suficiente), excluindo '
     'feriados. Programação de trabalho fora desse horário comercial normal ("Trabalho fora do horário '
     'comercial") requer considerações sobre o pessoal e precisa ser planejada com pelo menos 20 (vinte) dias '
     'úteis de antecedência da necessidade.')

h2('5.5 Alocação dos Recursos SFDC')
para('Os recursos dos Serviços Profissionais serão alocados de acordo com um plano de alocação mutuamente '
     'acordado no início do engajamento. Desde que envie uma notificação prévia com 2 (duas) semanas de '
     'antecedência à SFDC, o Cliente poderá solicitar alterações ao plano de alocação. Caso o Cliente solicite '
     'alterações com notificação prévia inferior a 2 (duas) semanas, a continuidade dos recursos da SFDC não '
     'poderá ser garantida e a SFDC poderá faturar o Cliente pelas horas planejadas e quaisquer custos não '
     'reembolsáveis (ex.: T&E) relativos àquelas 2 (duas) semanas. Exemplos de alterações incluem, sem '
     'limitação: mudança da data de início do projeto; atrasos ou suspensão das atividades do projeto; '
     'redução da alocação de recursos; e extensão desta SOW além da data final do engajamento.')

h2('5.6 Consentimento para o Uso de Subcontratados e Subprocessadores Aprovados')
para('A SFDC será responsável pela execução dos Serviços Profissionais por seu pessoal, incluindo '
     'subcontratados (se houver), exceto conforme de outra forma aqui especificado. O Cliente concorda, ao '
     'assinar esta SOW, que a SFDC poderá utilizar subcontratados para a execução dos Serviços Profissionais. '
     'Caso o Cliente forneça à SFDC acesso a Dados Pessoais, o Cliente concorda que o pessoal da SFDC e suas '
     'Afiliadas poderão processar tais Dados Pessoais mediante o uso de ferramentas de terceiros baseadas na '
     'nuvem listadas na documentação de subprocessadores da Salesforce a partir da Data de Vigência. Para os '
     'fins desta SOW, Dados Pessoais significam informações eletrônicas relativas a uma pessoa física '
     'identificada ou identificável, fornecidas à SFDC para a execução dos Serviços Profissionais.')

doc.add_page_break()

# ============================================================
# SEÇÃO 6 — RESPONSABILIDADES DO CLIENTE
# ============================================================
h1('6. Responsabilidades do Cliente')
para('A SFDC espera executar as atividades descritas nesta SOW como um fluxo ininterrupto de trabalho. O '
     'desempenho tempestivo dos Serviços Profissionais requer colaboração contínua entre a SFDC e o Cliente. '
     'O Cliente é responsável por determinadas tarefas importantes, contribuições e revisões em tempo hábil '
     'para permitir que a SFDC execute suas obrigações. Atrasos causados pelo Cliente podem resultar em '
     'encargos adicionais pelo tempo do recurso e a SFDC não será responsável por atrasos ou danos decorrentes '
     'do descumprimento do Cliente de suas obrigações.')

h2('6.1 Cooperação do Cliente')
para('O Cliente cooperará com a SFDC na execução dos Serviços Profissionais, incluindo, sem limitação:')
for x in [
    'Determinar prazos e cronogramas de recursos mutuamente aceitáveis com a equipe de implementação da SFDC.',
    'Entregar em tempo hábil as garantias e demais obrigações do Cliente exigidas por esta SOW.',
    'Responder prontamente às consultas da SFDC e fornecer informações, dados e feedback completos, precisos e tempestivos.',
    'Colaborar para mitigar riscos, escalar e resolver problemas rapidamente e priorizar esforços para cumprir os prazos acordados.',
    'Participar de reuniões semanais de status (Gerente de Projeto e Proprietário do Produto do Cliente + Gerente de Projeto/Engajamento da SFDC).',
    'Participar de reunião mensal do Comitê de Direção (Steering Committee) entre os executivos responsáveis e gerentes de projeto de ambas as partes.',
]:
    bullet(x)

h2('6.2 Equipe do Cliente')
para('O Cliente disponibilizará recursos adequados e com conhecimento suficiente da Aplicação, suas '
     'funcionalidades padrão e terminologia, para participação ativa e contínua (revisão, feedback, '
     'aprovações e tomada de decisão em tempo hábil), incluindo os papéis abaixo:')
roles = [
    ('Executivo Responsável', 'Liderança e orientação geral, participação no Comitê de Direção e ponto de escalação.'),
    ('Gerente do Projeto', 'Ponto único de contato com a SFDC; co-gerencia plano e cronograma; coordena recursos, decisões e aceitação.'),
    ('Proprietário do Produto', 'Responsável pela visão do produto e priorização do backlog; fornece conteúdo para Histórias de Usuário e aceita cada História após demonstração. Na Claro: Product Owners com autonomia de aprovação por LOB (Móvel, TV, Banda Larga, Fixo e Aparelhos).'),
    ('Especialistas em Assuntos de Negócios (SMEs)', 'Definem requisitos de negócio e casos de teste; responsáveis por UAT e testes de usabilidade.'),
    ('Líder Funcional', 'Define requisitos de sistema e casos de teste; supervisiona a transferência de conhecimento.'),
    ('Líder de Tecnologia/TI', 'Suporte de configuração de rede, requisitos de segurança e SSO, decisões de arquitetura e acesso aos sistemas. Na Claro: interlocução com Billing, OSS/BSS, MuleSoft, Okta e SailPoint.'),
    ('Líder de Gestão de Alteração Organizacional', 'Análise de stakeholders, plano de comunicação, adoção pelos usuários e captura do valor de negócio.'),
    ('Líder de Treinamento', 'Estratégia, material e administração do treinamento dos usuários finais.'),
]
for r_name, r_desc in roles:
    p = doc.add_paragraph(style='List Bullet')
    rr = p.add_run(f'{r_name}: '); rr.bold = True; rr.font.color.rgb = SF_BLUE
    p.add_run(r_desc)

h2('6.3 Obrigações Adicionais')
h3('Acesso à Tecnologia e à Aplicação')
para('O Cliente adquirirá, instalará, hospedará, testará, implantará, monitorará e manterá os equipamentos, '
     'infraestrutura, software, Internet estável com largura de banda suficiente e demais tecnologias '
     'necessárias à prestação dos Serviços Profissionais, habilitando as permissões de acesso do pessoal da '
     'SFDC à Aplicação conforme razoavelmente necessário. A SFDC terá acesso apenas à Aplicação do Cliente, e '
     'não aos sistemas não-SFDC, salvo acordo mútuo. O acesso do pessoal da SFDC será revogado na conclusão '
     'do Prazo da SOW.')
h3('Obrigações específicas do engajamento (Claro)')
for x in [
    'Provisionar a Sandbox funcional da Communications Cloud antes do início da Fase 3 (PoC Build) — pré-requisito inegociável.',
    'Fornecer amostras das matrizes de elegibilidade e precificação vigentes (planilhas ou diagramas lógicos).',
    'Fornecer documentação de APIs dos sistemas legados (Billing, OSS/BSS) e do georreferenciamento por CEP, quando aplicável.',
    'Prover acessos, ambientes e credenciais necessários em tempo hábil.',
    'Revisar e aprovar entregáveis nos prazos de aceitação acordados (Seção 3.3).',
]:
    bullet(x)
h3('Disponibilidade e Qualidade de Dados')
para('O Cliente é responsável pela disponibilidade e qualidade dos dados a serem utilizados/integrados, '
     'incluindo transformação, limpeza, agregação, desduplicação e testes de garantia de qualidade antes de '
     'submetê-los à SFDC, bem como pelo acesso às fontes de dados e endpoints necessários.')
h3('Uso dos Serviços')
para('O Cliente é responsável por seu uso e distribuição dos Entregáveis resultantes dos Serviços Profissionais.')

h2('6.4 Produtos e Serviços de Terceiros')
para('"Produtos de Terceiros" significa produtos ou serviços fornecidos ou adquiridos pelo Cliente em conexão '
     'com os Serviços Profissionais (ex.: sistemas de Billing, OSS/BSS, MuleSoft, IdP Okta, SailPoint, e '
     'aplicações de um Marketplace Salesforce como o AppExchange). A SFDC não é responsável nem garante '
     'Produtos de Terceiros. O Cliente é responsável (i) pela aquisição, licenciamento e uso; (ii) salvo se '
     'definido como escopo na Seção 2, pela integração, instalação, implementação e configuração; (iii) pela '
     'coordenação e gestão dos respectivos fornecedores; e (iv) por assegurar que estejam prontos em tempo '
     'razoável para a execução dos Serviços Profissionais.')

doc.add_page_break()

# ============================================================
# SEÇÃO 7 — TERMOS GERAIS
# ============================================================
h1('7. Termos Gerais')

h2('7.1 Controle de Mudança')
para('No caso de mudanças no escopo do trabalho ou em outros termos desta SOW, as partes trabalharão juntas '
     'de boa-fé para acordar um Pedido de Alteração apropriado, seguindo o Processo de Controle de Mudança '
     'definido abaixo. A SFDC não terá obrigação de executar serviços profissionais adicionais ou modificados '
     'na ausência de acordo sobre um Pedido de Alteração. As taxas ou o Prazo da SOW podem estar sujeitos a '
     'Pedido de Alteração em caso de mudança material ou deficiência nas informações fornecidas pelo Cliente, '
     'evento imprevisto que altere materialmente as necessidades de serviço, mudança na lei/regulamentação '
     'aplicável, ou evento de força maior.')
h3('Definições de Controle de Mudança')
for term, d in [
    ('Controle de Mudanças', 'descreve o processo para gerenciar mudanças potenciais do projeto.'),
    ('Solicitação de Mudança', 'é uma solicitação informal do Cliente ou da SFDC para modificar os Serviços Profissionais previstos nesta SOW.'),
    ('Pedido de Alteração', 'é uma emenda formal a esta SOW, conforme definido na Seção 1.4.'),
]:
    p = doc.add_paragraph(style='List Bullet')
    rr = p.add_run(f'"{term}" '); rr.bold = True; rr.font.color.rgb = SF_BLUE
    p.add_run(d)
h3('Processo de Controle de Mudança')
para('A SFDC continuará executando os Serviços Profissionais de acordo com esta SOW até que as partes '
     'concordem por escrito sobre a alteração. Resumo do processo:')
for x in [
    'O Cliente ou a SFDC apresenta uma Solicitação de Mudança por escrito (e-mail aceitável), com descrição, justificativa e impacto sobre os Serviços Profissionais.',
    'Ambas as partes revisam o mérito e o impacto sobre escopo, recursos, cronograma, honorários e demais termos.',
    'As partes decidem aceitar ou rejeitar a alteração solicitada.',
    'Se aceita, a SFDC prepara um Pedido de Alteração descrevendo as alterações aplicáveis.',
    'A SFDC assina e envia o Pedido de Alteração ao Cliente.',
    'O Cliente assina o Pedido de Alteração e, se aplicável, gera a PO associada.',
    'Os Serviços Profissionais do Pedido de Alteração totalmente assinado passam a integrar esta SOW.',
]:
    bullet(x)

h2('7.2 Segmentação')
para('O Cliente reconhece que esta SOW é limitada aos Serviços Profissionais e não transmite qualquer direito '
     'de uso da Aplicação, cujo uso será regido por acordo separado. O Cliente concorda que sua aquisição dos '
     'Serviços Profissionais não está condicionada à entrega de qualquer funcionalidade ou característica '
     'futura da Aplicação (exceto Entregáveis sob esta SOW), nem a comentários públicos, orais ou escritos, da '
     'SFDC a respeito de funcionalidades futuras.')

h2('7.3 Rescisão')
para('O Cliente poderá rescindir esta SOW mediante notificação por escrito com antecedência de 30 (trinta) '
     'dias. Qualquer uma das partes pode rescindir por justa causa: (i) mediante notificação escrita com '
     'trinta (30) dias de antecedência de uma violação material, se não sanada ao término do período; ou (ii) '
     'se a outra parte se tornar objeto de pedido de falência, insolvência, recuperação judicial, liquidação '
     'ou cessão em benefício de credores. Serviços prestados e despesas incorridas até a data efetiva da '
     'rescisão são devidos à SFDC.')

h2('7.4 Geral')
para('Esta SOW poderá ser assinada em vias, cada uma considerada um original, e está sujeita aos termos e '
     'condições do Contrato, exceto se de outra forma expressamente previsto nesta SOW. No caso de conflito '
     'entre qualquer termo desta SOW e o Contrato, os termos desta SOW prevalecerão. EM TESTEMUNHO DO QUE, as '
     'partes assinam esta SOW por seus respectivos representantes devidamente autorizados, conforme '
     'identificados abaixo.')

doc.add_paragraph()
doc.add_paragraph()

# Assinaturas
h2('Aceite e Assinaturas')
sig = doc.add_table(rows=2, cols=2)
sig.style = 'Table Grid'
labels = [
    ('CLARO BRASIL', 'SALESFORCE (SFDC)'),
]
c = sig.rows[0].cells
c[0].text = ''; c[1].text = ''
c[0].paragraphs[0].add_run('CLARO BRASIL').bold = True
c[1].paragraphs[0].add_run('SALESFORCE (SFDC)').bold = True
c2 = sig.rows[1].cells
for cell in c2:
    cell.text = ''
    cell.add_paragraph('\n\nNome: ______________________________')
    cell.add_paragraph('Cargo: _____________________________')
    cell.add_paragraph('Data: ______________________________')
    cell.add_paragraph('Assinatura: _________________________')

# Rodapé
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.text = 'Salesforce Professional Services — Confidencial · CLARO Brasil CPQ Design · SOW v1.0 (Rascunho)'
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in fp.runs:
    run.font.size = Pt(8); run.font.color.rgb = GREY

out = '/Users/nfilho/claude/claro-cpq-design-sow/drafts/SOW_CLARO_CPQ_Design_v1.docx'
doc.save(out)
print('SOW gerado:', out)

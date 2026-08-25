# -*- coding: utf-8 -*-
"""Edita o template IronClad (SOW #05719988) preenchendo placeholders com o
escopo Claro CPQ Design conforme ROM final. Preserva dados reais, tabelas de
reforma tributária e tags de assinatura IronClad."""
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

SRC = "sources/IRONCLAD_Claro_SOW_template.docx"
OUT = "drafts/SOW_CLARO_CPQ_Design_v2_IronClad.docx"


def set_text(p_elem, text):
    """Substitui o texto do parágrafo mantendo o rPr do primeiro run."""
    runs = p_elem.findall(qn("w:r"))
    if not runs:
        r = p_elem.makeelement(qn("w:r"), {})
        t = r.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
        t.text = text
        r.append(t)
        p_elem.append(r)
        return
    first = runs[0]
    ts = first.findall(qn("w:t"))
    if ts:
        ts[0].text = text
        ts[0].set(qn("xml:space"), "preserve")
        for extra in ts[1:]:
            first.remove(extra)
    else:
        t = first.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
        t.text = text
        first.append(t)
    for r in runs[1:]:
        p_elem.remove(r)


def clone_after(ref_elem, text):
    """Clona ref_elem (preserva estilo/numeração), troca o texto e insere após."""
    new = deepcopy(ref_elem)
    set_text(new, text)
    ref_elem.addnext(new)
    return new


def set_cell(cell, text):
    """Define o texto da célula no primeiro parágrafo, removendo os demais."""
    paras = cell.paragraphs
    set_text(paras[0]._p, text)
    for p in paras[1:]:
        p._p.getparent().remove(p._p)


d = Document(SRC)
P = d.paragraphs

# ---------------------------------------------------------------- Objetivos
OBJ = [
    "Reduzir o time-to-market (TTM) de novas ofertas, planos e promoções, "
    "substituindo matrizes e customizações legadas por capacidades declarativas "
    "nativas (EPC e Motor de Regras) da Communications Cloud.",
    "Unificar o catálogo comercial das cinco linhas de negócio (Móvel, TV, Banda "
    "Larga, Fixo e Aparelhos) em um modelo de dados único, aderente ao padrão TM Forum SID.",
    "Habilitar a venda ágil de combos multi-LOB, com cross-sell e up-sell, por meio "
    "do carrinho de vendas (CPQ/Cart).",
    "Modelar de forma declarativa as regras de elegibilidade, compatibilidade e "
    "precificação, reduzindo a dependência de código customizado.",
    "Validar, por meio de uma Prova de Conceito (PoC), o padrão de decomposição de "
    "pedidos (Order Management) e o payload de integração com os sistemas legados.",
]
set_text(P[25]._p, OBJ[0])
ref = P[25]._p
for o in OBJ[1:]:
    ref = clone_after(ref, o)

# --------------------------------------------------- EPC / Product Catalog
set_text(P[32]._p,
    "A SFDC desenhará e modelará o Enterprise Product Catalog (EPC) da Communications "
    "Cloud como catálogo comercial principal, contemplando até 100 (cem) produtos "
    "comerciais (até 20 por linha de negócio: Móvel, TV, Banda Larga, Fixo e Aparelhos), "
    "organizados em duas camadas — Produto Comercial e Produto Técnico/CFS —, conforme o "
    "padrão TM Forum SID. O escopo inclui a definição do product template (Object Type / "
    "Product Specification), a estrutura de Catalogs e a modelagem de até 5 (cinco) "
    "promoções/descontos. A configuração efetiva no EPC será realizada em escopo de Prova "
    "de Conceito (PoC): 1 (uma) oferta por linha de negócio precificada e publicada, "
    "usando a Comms Cloud como catálogo comercial principal.")

set_text(P[34]._p,
    "Desenvolvimento: O Cliente ou o Integrador de Sistemas (SI) é responsável pelas "
    "atividades de DevOps. O Salesforce Industries utilizará sempre a capacidade nativa do "
    "produto disponível em disponibilidade geral (GA), em vez de opções personalizadas. As "
    "integrações com sistemas externos (por exemplo, verificação de crédito, "
    "cobertura/serviceability, inventário e pagamentos) não fazem parte do escopo desta PoC "
    "e, quando aplicável, serão simuladas/mockadas em ambiente Sandbox.")

# --------------------------------------------------------------- CPQ / Cart
set_text(P[41]._p,
    "Desenvolvimento: O Cliente ou o Integrador de Sistemas (SI) é responsável pelas "
    "atividades de DevOps. O Salesforce Industries utilizará sempre a capacidade nativa do "
    "produto disponível em disponibilidade geral (GA), em vez de opções personalizadas. "
    "Nesta PoC, as interfaces com sistemas externos do fluxo de Cotação para Pedido serão "
    "simuladas/mockadas; a construção das integrações produtivas é responsabilidade do "
    "Cliente ou do SI.")

# ---------------------------------------- Deployment / Documentação (era EN)
set_text(P[43]._p,
    "Implantação em Produção e Go-live. A implantação em ambiente de Produção e o go-live "
    "não fazem parte do escopo desta SOW. Os Serviços Profissionais compreendem o desenho, "
    "a modelagem e a validação de uma Prova de Conceito (PoC) em ambiente Sandbox. A "
    "eventual industrialização e implantação em Produção serão objeto de SOW específica.")
set_text(P[44]._p,
    "Documentação. A SFDC fornecerá os artefatos de desenho (Detalhamento Funcional das "
    "Ofertas, Product Model Blueprint e recomendações de regras de negócio) e o Relatório "
    "Final, que refletem o estado da PoC configurada.")

# ------------------------------------------------ Premissas de Escopo (add)
PREM = [
    "A entrega do desenvolvimento em 8 (oito) semanas está condicionada ao aceite, pelo "
    "Cliente, da Cláusula de Consentimento para o Uso de Ferramentas de Inteligência "
    "Artificial (IA) desta SOW. Sem tal aceite, o cronograma poderá ser reajustado "
    "mediante Pedido de Alteração.",
    "O Cliente disponibilizará ambiente Sandbox adequado e os acessos necessários no "
    "início do projeto.",
    "O escopo de configuração está limitado a até 100 produtos comerciais (até 20 por "
    "linha de negócio) e até 5 promoções, com 1 oferta por linha de negócio publicada na PoC.",
    "O Cliente disponibilizará, em tempo hábil, os especialistas (SMEs) de TI e de Negócio "
    "e a documentação das ofertas atuais.",
    "As integrações com sistemas externos serão simuladas/mockadas em Sandbox; nenhuma "
    "integração produtiva será construída neste escopo.",
]
ref = P[51]._p
for x in PREM:
    ref = clone_after(ref, x)

# ------------------------------------------------ Exclusões de Escopo (add)
EXCL = [
    "Implantação em ambiente de Produção e go-live",
    "Construção de integrações produtivas com sistemas legados (Billing, OSS/BSS, CRM e "
    "meios de pagamento)",
    "Desenvolvimento de código customizado (Apex, LWC) além do estritamente necessário à "
    "PoC, a critério da SFDC",
    "Migração, limpeza ou deduplicação de dados históricos",
    "Redesenho de processos de orquestração e fulfillment",
    "Fornecimento de licenças de software Salesforce ou de terceiros",
    "Serviços gerenciados (AMS) e hypercare pós go-live",
]
ref = P[61]._p
for x in EXCL:
    ref = clone_after(ref, x)

# ---------------------------------------------------------------- Cronograma
set_text(P[64]._p,
    "A SFDC estima que o projeto tenha duração de 12 (doze) semanas. Esta estimativa "
    "baseia-se na suposição da SFDC de que será capaz de executar os Serviços Profissionais "
    "sem atrasos iniciados pelo Cliente. Um cronograma preliminar é estabelecido abaixo, e "
    "um cronograma mais detalhado do projeto é produzido durante a fase de Definição do "
    "projeto. Os Serviços Profissionais começam em uma data de início a ser acordada "
    "mutuamente por escrito (e-mail aceitável), mas não antes da Data Efetiva desta SOW. O "
    "desenvolvimento (Marcos 1 a 3) é entregue em 8 (oito) semanas, mediante o uso de "
    "ferramentas de Inteligência Artificial (IA), conforme consentimento do Cliente; as "
    "semanas 9 a 12 (Marco 4) destinam-se ao acompanhamento do Teste de Aceitação do "
    "Usuário (UAT), à transferência de conhecimento (KT) e à consolidação do Relatório Final.")

# ---------------------------------------------------------- Tabela de Marcos
# As colunas "Semana" (3 e 4) são uma célula mesclada: usamos um intervalo.
t = d.tables[0]
set_cell(t.rows[0].cells[3], "Semana Estimada (Início–Fim)")

CRIT = {
    1: ("Atividade: Workshops de descoberta com TI e Negócio para mapeamento e "
        "detalhamento funcional das ofertas, planos e serviços atuais das cinco linhas de "
        "negócio. Critério de aceite: Documento de Detalhamento Funcional das Ofertas "
        "revisado e aceito formalmente pelo Cliente.", "1–2"),
    2: ("Atividade: Desenho do modelo de dados no EPC (camadas comercial e técnica/CFS, "
        "Object Types e atributos) e especificação das regras de elegibilidade, "
        "compatibilidade e precificação no Motor de Regras. Critério de aceite: Product "
        "Model Blueprint (PMB) e recomendações de regras de negócio aprovados pelo Cliente.", "3–4"),
    3: ("Atividade: Configuração da Prova de Conceito em Sandbox — 1 oferta por linha de "
        "negócio, contemplando prateleira, carrinho (CPQ), validações e payload de "
        "decomposição de pedido (Order Management). Critério de aceite: PoC demonstrada e "
        "validada em Sandbox, com UAT aceito pelo Cliente. Desenvolvimento entregue em 8 "
        "semanas mediante uso de ferramentas de IA (ver Cláusula de Consentimento de IA).", "5–8"),
    4: ("Atividade: Acompanhamento do UAT do Cliente, workshops de transferência de "
        "conhecimento (TI e Negócio) e consolidação do Relatório Final. Critério de aceite: "
        "Conclusão dos workshops de KT e entrega e aceite do Relatório Final do projeto.", "9–12"),
}
for r, (crit, semana) in CRIT.items():
    set_cell(t.rows[r].cells[2], crit)
    set_cell(t.rows[r].cells[3], semana)

# ------------------------------------------- Cláusula de Consentimento de IA
# Inserida logo após a cláusula de subprocessadores (para 143).
ai_title = deepcopy(P[142]._p)   # Heading 2
set_text(ai_title, "Consentimento para o Uso de Ferramentas de Inteligência Artificial (IA)")
P[143]._p.addnext(ai_title)

ai_body = deepcopy(P[143]._p)    # Heading 3 numerado
set_text(ai_body,
    "O Cliente reconhece e concorda, ao assinar esta SOW, que a SFDC e suas Afiliadas "
    "poderão utilizar ferramentas de Inteligência Artificial (IA), incluindo IA "
    "generativa, próprias ou de subprocessadores aprovados, para apoiar a execução dos "
    "Serviços Profissionais — por exemplo, na aceleração de descoberta, modelagem de dados, "
    "geração e revisão de configuração, documentação e artefatos de projeto. O uso de tais "
    "ferramentas observará os princípios de IA Confiável (Trusted AI) da Salesforce, os "
    "termos do Contrato e as políticas aplicáveis de proteção de dados, sendo a relação de "
    "subprocessadores e ferramentas de terceiros aprovados aquela disponível no endereço "
    "referido na cláusula anterior. O prazo de entrega do desenvolvimento em 8 (oito) "
    "semanas e os Honorários estabelecidos nesta SOW estão condicionados a este "
    "consentimento; na ausência dele, o cronograma, os Marcos e os Honorários poderão ser "
    "reajustados mediante Pedido de Alteração.")
ai_title.addnext(ai_body)

d.save(OUT)
print("Salvo:", OUT)
